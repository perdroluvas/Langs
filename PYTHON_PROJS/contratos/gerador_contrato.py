from datetime import datetime, timedelta
from num2words import num2words
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

def number_to_words_pt(number):
    """Convert number to words in Portuguese"""
    return num2words(number, lang='pt_BR').upper()

def format_currency(value):
    """Format currency value in Brazilian Real"""
    return f"R$ {value:.2f}"

def add_centered_paragraph(document, text, bold=False, space_after=True):
    """Add centered paragraph with optional bold and spacing"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    if space_after:
        paragraph.space_after = Pt(12)

def generate_contract():
    print("\n=== GERADOR DE CONTRATO DE LOCAÇÃO ===\n")
    
    # Get user input
    locador_nome = input("Sandra ou Janja?")
    locatario_nome = input("Nome completo do locatário: ").upper()
    locatario_rg = input("RG do locatário (com órgão expedidor): ")
    locatario_cpf = input("CPF do locatário (formato XXX.XXX.XXX-XX): ")
    
    valor_aluguel = float(input("Valor do aluguel (R$): "))
    valor_aluguel_extenso = number_to_words_pt(valor_aluguel)
    
    dia_vencimento = int(input("Dia do vencimento (1-31): "))
    dia_vencimento_extenso = number_to_words_pt(dia_vencimento)
    
    periodo = input("Período do contrato (6 para SEIS MESES ou 12 para UM ANO): ")
    if periodo == "6":
        periodo_extenso = "SEIS MESES"
        periodo_numerico = "06 (SEIS)"
    else:
        periodo_extenso = "UM ANO"
        periodo_numerico = "01 (UM)"
    
    # Calculate dates
    data_inicio = datetime.now()
    if periodo == "6":
        data_fim = data_inicio + timedelta(days=180)
    else:
        data_fim = data_inicio + timedelta(days=365)

    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_centered_paragraph(doc, "CONTRATO DE LOCAÇÃO DE IMÓVEL RESIDENCIAL", bold=True)

    # Contract body
    if locador_nome == "Sandra":
    # Dados de Elissandra Batista dos Santos (LOCADOR)
        doc.add_paragraph(f"""Pelo presente instrumento particular de locação, de um lado, ELISSANDRA BATISTA DOS SANTOS, brasileira, solteira, portador da C.I Nº2395836 e CPF: Nº 429.041.962-34, residente na TRV Timbó n 60 – Pedreira- –Belém – PA, denominado LOCADOR, e de outro, o Sr(a). {locatario_nome}, portador da RG Nº {locatario_rg} e CPF Nº {locatario_cpf}, doravante denominado LOCATÁRIO, têm, entre si, justos e contratados, o presente contrato, consoante as cláusulas e condições abaixo:""")

    elif locador_nome == "Janja":
    # Dados de Elisangela Batista dos Santos (LOCADOR)
        doc.add_paragraph(f"""Pelo presente instrumento particular de locação, de um lado, ELISSANGELA BATISTA DOS SANTOS, brasileira, divorciada, portador da C.I Nº2518063 e CPF: Nº 352.202.952-68, residente na TRV Vileta n 2501 – Marco- –Belém – PA, denominado LOCADOR, e de outro, o Sr(a). {locatario_nome}, portador da RG Nº {locatario_rg} e CPF Nº {locatario_cpf}, doravante denominado LOCATÁRIO, têm, entre si, justos e contratados, o presente contrato, consoante as cláusulas e condições abaixo:""")

    doc.add_paragraph("Constitui objeto do presente Contrato a locação do imóvel localizado no Conjunto Cidade Nova III, We 18, n° 22 Apartamento 3, Bairro do Coqueiro –Ananindeua-PA, CEP: 67130675")

    # Cláusula Segunda
    add_centered_paragraph(doc, "CLÁUSULA SEGUNDA", bold=True)
    
    doc.add_paragraph(f"""O prazo de locação é de {periodo_numerico} {periodo_extenso}, com termo inicial em ({data_inicio.strftime('%d/%m/%Y')}) e termo final em ({data_fim.strftime('%d/%m/%Y')}), data em que o LOCATÁRIO se obriga a restituir o imóvel livre e desocupado, em condições idênticas à que recebeu, ressalvando o desgaste natural do imóvel, independentemente de aviso ou notificação.""")

    # [Continue with all other clauses from the original contract...]
    # Add all other clauses here...

    add_centered_paragraph(doc, "PARÁGRAFO PRIMEIRO -", bold = True)

    doc.add_paragraph(f"""Findo o prazo estipulado no caput desta Cláusula, operar-se-á o término da avença, sendo que eventual prorrogação tão somente ocorrerá por meio de adiamento contratual, de acordo com a conveniência das partes.""")

    add_centered_paragraph(doc, "PARÁGRAFO SEGUNDO -", bold = True)

    doc.add_paragraph(f"""As Benfeitorias que o LOCADOR deixou no imóvel encontram-se em perfeitas condições de uso, e integram este instrumento para todos os fins e efeitos de direito.""")




    # Cláusula Terceira
    add_centered_paragraph(doc, "CLÁUSULA TERCEIRA", bold=True)
    
    doc.add_paragraph(f"""O aluguel mensal fica estipulado em {format_currency(valor_aluguel)} ({valor_aluguel_extenso}), devendo ser pago pelo LOCATÁRIO, até o dia {dia_vencimento} ({dia_vencimento_extenso}) de cada mês, diretamente ao LOCADOR.""")


    # [Add all remaining clauses...]


    add_centered_paragraph(doc, "PARÁGRAFO PRIMEIRO -", bold = True)

    doc.add_paragraph(f"""O valor locativo será reajustado anualmente, no caso de prorrogação do presente contrato, de acordo com a variação acumulada do INCC ou, se extinto, pelo IGPM/FGV. Na ausência destes índices será eleito, aquele que venha a substituí-los.""")


    add_centered_paragraph(doc, "PARÁGRAFO SEGUNDO -", bold = True)
    
    doc.add_paragraph(f"""Sobre o aluguel pago após o quinto dia do respectivo vencimento, incidirá multa moratória de R$5,00 (cinco reais) AO DIA. além das despesas contratuais e extras que os locadores despenderem para a ressalva de seus direitos. Na hipótese de o atraso ultrapassar 30 (trinta) dias, será o débito acrescido de juros de 10% (dez por cento) ao mês e de correção monetária, com base na variação dos incides aludidos no parágrafo segundo desta cláusula.""")
    
    add_centered_paragraph(doc, "CLÁUSULA QUARTA", bold=True)
    doc.add_paragraph("O LOCADOR deixa reservado seu direito de receber qualquer aluguel fora do prazo contratado, "
                  "sem que isso importe em novação deste Contrato. Qualquer despesa judicial ou extrajudicial, "
                  "feita pelo LOCADOR para a cobrança de aluguéis, fora do prazo previsto, inclusive honorários "
                  "de advogado, correrá por conta do LOCATÁRIO e deverá ser paga juntamente com o aluguel devido.")

    add_centered_paragraph(doc, "CLÁUSULA QUINTA", bold=True)
    doc.add_paragraph("O imóvel deste Contrato destina-se exclusivamente para fins de ser aluguel residencial.")

    add_centered_paragraph(doc, "CLÁUSULA SEXTA", bold=True)
    doc.add_paragraph("Fica vedada a sublocação do imóvel ou a cessão dos direitos decorrentes deste instrumento a "
                  "terceiros, mesmo que parcial ou temporária, seja a que título for, por parte do LOCATÁRIO, "
                  "sem a expressa anuência do LOCADOR.")

    add_centered_paragraph(doc, "CLÁUSULA SÉTIMA", bold=True)
    doc.add_paragraph("A LOCATÁRIA declara que vistoriou o imóvel, objeto deste Contrato e que tem pleno conhecimento "
                  "de que está ele em perfeitas condições de uso para a finalidade prevista na Cláusula Quinta.")

    add_centered_paragraph(doc, "CLÁUSULA OITAVA", bold=True)
    doc.add_paragraph("Além do aluguel mensal, incumbirá O LOCATÁRIO o pagamento da fatura de energia elétrica, ficando "
                  "por conta do LOCADOR, as despesas incidentes sobre o imóvel, com, por exemplo, água, Imposto "
                  "Predial e Territorial Urbano (IPTU).")

    add_centered_paragraph(doc, "PARÁGRAFO PRIMEIRO", bold=True)
    doc.add_paragraph("Os encargos da locação, especificados no caput desta cláusula, são de inteira responsabilidade do "
                  "LOCATÁRIO, que se obriga a pagá-los em seus respectivos vencimentos, devendo comprová-los ao "
                  "LOCADOR sempre que solicitado, e, em especial, quando do encerramento do Contrato.")

    add_centered_paragraph(doc, "PARÁGRAFO SEGUNDO", bold=True)
    doc.add_paragraph("O LOCATÁRIO restituirá o imóvel locado, nas mesmas condições as quais o recebeu, quais sejam, "
                  "pintado com tinta látex, na cor em que foi entregue, sendo que as instalações elétricas, hidráulicas "
                  "e acessórios deverão também, estar em perfeitas condições de funcionamento.")

    add_centered_paragraph(doc, "PARÁGRAFO TERCEIRO", bold=True)
    doc.add_paragraph("Durante vigência deste contrato, o LOCADOR fica autorizado a transferir para o nome do LOCATÁRIO, "
                  "as faturas de energia elétrica da C/C:_____________________ , devidamente instalada no imóvel locado.")

    add_centered_paragraph(doc, "CLÁUSULA NONA", bold=True)
    doc.add_paragraph("Quaisquer benfeitorias a serem introduzidas internamente ou externamente no imóvel dependerão de "
                  "prévia anuência do LOCADOR, as quais, se efetivadas, se incorporarão ao imóvel, não cabendo nenhuma "
                  "indenização ao LOCATÁRIO, por parte do LOCADOR.")

    add_centered_paragraph(doc, "CLÁUSULA DÉCIMA", bold=True)
    doc.add_paragraph("Este contrato fica facultado o seu registro no Cartório de Registro Imobiliário competente, ademais, "
                  "o presente contrato passa a vigorar entre as partes a partir da assinatura do mesmo, as quais elegem "
                  "o foro da cidade de Ananindeua (PA) ou da cidade de Belém (PA), onde se situa o imóvel, para dirimirem "
                  "quaisquer dúvidas provenientes da execução e cumprimento do mesmo. Por estarem assim justos e contratados, "
                  "firmam o presente instrumento, em duas vias de igual teor, sem emenda ou rasuras, juntamente com a assinatura "
                  "de 2 (duas) testemunhas, para que surta seus legais e efeitos jurídicos.")

# Signature section
    doc.add_paragraph("\n\n")
    add_centered_paragraph(doc, f"ANANINDEUA, _____ de ________________ de {data_inicio.year}")
    
    doc.add_paragraph("\n\n")
    add_centered_paragraph(doc, "_" * 50)
    add_centered_paragraph(doc, "ELISANGELA BATISTA DOS SANTOS")
    add_centered_paragraph(doc, "Locador")
    
    doc.add_paragraph("\n")
    add_centered_paragraph(doc, "_" * 50)
    add_centered_paragraph(doc, locatario_nome)
    add_centered_paragraph(doc, "Locatário")

    # Save document
    filename = f"contrato_locacao_{locatario_nome.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    
    print(f"\nContrato gerado com sucesso! Arquivo salvo como: {filename}")

if __name__ == "__main__":
    try:
        generate_contract()
    except Exception as e:
        print(f"Erro ao gerar contrato: {str(e)}")
