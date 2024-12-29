from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import locale

def criar_procuracao():
    # Configurar localização para português
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
    
    # Solicitar dados do outorgante
    pessoa = input("Sandra ou Janja?")
    nome = input("Digite o nome completo do outorgante: ")
    nacionalidade = input("Digite a nacionalidade (brasileiro/brasileira): ")
    estado_civil = input("Digite o estado civil: ")
    cpf = input("Digite o CPF (apenas números): ")
    rg = input("Digite o RG: ")
    apartamento = input("Digite o número do apartamento (1-4): ")
    
    # Criar novo documento
    doc = Document()
    
    # Configurar estilo do título
    titulo = doc.add_paragraph("PROCURAÇÃO PARTICULAR")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].bold = True
    
    # Adicionar outorgado (informação constante)
    outorgado = doc.add_paragraph()
    outorgado.add_run("OUTORGADO: ").bold = True
    if pessoa == "Sandra":
        outorgado.add_run("Elissandra Batista dos Santos, brasileira, solteira, inscrita no CPF/MF sob o n° 429.041.962-34 e RG 2395836 PC/PA, residente e domiciliada sito à Travessa Timbó, passagem Gama Malcher 60, CEP 66085-390 - Belém, Pará.")
    elif pessoa == "Janja":
        outorgado.add_run("Elisangela Batista dos Santos, brasileira, divorciada, inscrita no CPF/MF sob o n° 352.202.952-68 e RG 2518063 PC/PA, residente na TRV Vileta n° 2501 – Marco – Belém – PA, CEP 66093-345 - Belém, Pará.")
    # Adicionar outorgante (informação variável)
    outorgante = doc.add_paragraph()
    outorgante.add_run("OUTORGANTE: ").bold = True
    outorgante.add_run(f"{nome}, {nacionalidade}, {estado_civil}, inscrito(a) no CPF/MF sob o n° {cpf} e RG {rg} SSP/PA residente Conjunto Cidade Nova IV, We 18, n° 32 Apartamento {apartamento}A, Bairro do Coqueiro –Ananindeua-PA, CEP: 67130-675, Belém, Pará.")
    
    # Adicionar poderes
    poderes = doc.add_paragraph()
    poderes.add_run("PODERES: ").bold = True
    poderes.add_run("Por este instrumento particular de Procuração, nomeio e constituo meu bastante Procurador acima, a quem confio amplos, gerais e ilimitado poderes para representar junto a empresa EQUATORIAL ENERGIA, a fim de solicitar os serviços de pedido de ligações, transferência de titularidade, pedido de baixas ou qualquer ato que se fizer necessário para o bem e fiel cumprimento do presente mandato. A presente tem validade até a conclusão do processo.")
    
    # Adicionar data atual
    data_atual = datetime.now()
    data = doc.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data.add_run(f"Belém-PA, {data_atual.day} de {data_atual.strftime('%B')} de {data_atual.year}")
    
    # Adicionar linha para assinatura
    doc.add_paragraph("\n\n")
    assinatura = doc.add_paragraph("_______________________________")
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    nome_assinatura = doc.add_paragraph(nome)
    nome_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    outorgante_assinatura = doc.add_paragraph("Outorgante")
    outorgante_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Gerar nome do arquivo
    nome_arquivo = f"{nome.lower().replace(' ', '_')}_outorga.docx"
    
    # Salvar documento
    doc.save(nome_arquivo)
    print(f"\nDocumento gerado com sucesso: {nome_arquivo}")

if __name__ == "__main__":
    try:
        criar_procuracao()
    except Exception as e:
        print(f"Erro ao gerar documento: {str(e)}")
